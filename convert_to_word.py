import re
from docx import Document
from docx.shared import Pt
import os

# Paths
md_path = "/Users/asap/.gemini/antigravity/brain/179de771-30be-4d97-8772-bf276ddfdc59/technical_report.md"
docx_path = "/Users/asap/.gemini/antigravity/brain/179de771-30be-4d97-8772-bf276ddfdc59/technical_report.docx"

def convert_md_to_docx(md_content, doc):
    lines = md_content.split('\n')
    
    in_list = False
    in_table = False
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines
        if not line:
            if in_list:
                in_list = False
            doc.add_paragraph()
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            continue
            
        # Lists
        if line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.', line):
            p = doc.add_paragraph(style='List Bullet' if not re.match(r'^\d+\.', line) else 'List Number')
            # Handle bold text in lists
            text = line[2:] if not re.match(r'^\d+\.', line) else re.sub(r'^\d+\.\s*', '', line)
            add_formatted_text(p, text)
            continue
            
        # Tables (Very basic support)
        if '|' in line:
            if not in_table:
                # This doesn't handle headers well but avoids crashing
                in_table = True
            # For now, just add as plain text to avoid table complexity in a quick script
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            continue
        else:
            in_table = False
            
        # Horizontal Rule
        if line == '---':
            doc.add_paragraph('_' * 30)
            continue
            
        # Normal paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line)

def add_formatted_text(paragraph, text):
    # Split by bold markers **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

try:
    if os.path.exists(md_path):
        with open(md_path, 'r') as f:
            md_content = f.read()
            
        document = Document()
        
        # Standard APA-ish formatting
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        convert_md_to_docx(md_content, document)
        document.save(docx_path)
        print(f"Successfully converted to {docx_path}")
    else:
        print(f"Error: {md_path} not found")
except Exception as e:
    print(f"Error during conversion: {e}")
